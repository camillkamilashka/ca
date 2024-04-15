import guietta
from docx import Document
from openpyxl import Workbook
from fabric_calculation import calculate_fabric_for_jacket, calculate_fabric_for_trousers, calculate_fabric_for_suit
from sewing_cost import calculate_sewing_cost

app = guietta.App("Clothing Calculator")

size_input = guietta.TextBox("Size")
fabric_cost_input = guietta.TextBox("Fabric Cost per meter")
fabric_label = guietta.Label("Fabric Needed for garment: ")
sewing_cost_label = guietta.Label("Sewing Cost: ")

def save_to_docx():
    doc = Document()
    doc.add_heading('Clothing Calculator Results', level=1)
    doc.add_paragraph(f"Size: {size_input.text()}")
    doc.add_paragraph(f"Fabric Cost per meter: {fabric_cost_input.text()}")
    doc.add_paragraph(fabric_label.text())
    doc.add_paragraph(sewing_cost_label.text())
    doc.save("clothing_calculator_results.docx")

def save_to_xls():
    wb = Workbook()
    sheet = wb.active
    sheet.title = "Results"
    sheet.append(["Size", size_input.text()])
    sheet.append(["Fabric Cost per meter", fabric_cost_input.text()])
    sheet.append(["", ""])
    sheet.append(["Garment", "Fabric Needed", "Sewing Cost"])
    # добавляем результаты расчетов
    wb.save("clothing_calculator_results.xlsx")

def calculate():
    size = float(size_input.text())
    fabric_cost = float(fabric_cost_input.text())

    fabric_needed_jacket = calculate_fabric_for_jacket(size)
    fabric_needed_trousers = calculate_fabric_for_trousers(size)
    fabric_needed_suit = calculate_fabric_for_suit(size)

    fabric_label.set(f"Fabric Needed for jacket: {fabric_needed_jacket}m\n"
                     f"Fabric Needed for trousers: {fabric_needed_trousers}m\n"
                     f"Fabric Needed for suit: {fabric_needed_suit}m")

    sewing_cost_jacket = calculate_sewing_cost(fabric_cost, fabric_needed_jacket)
    sewing_cost_trousers = calculate_sewing_cost(fabric_cost, fabric_needed_trousers)
    sewing_cost_suit = calculate_sewing_cost(fabric_cost, fabric_needed_suit)

    total_sewing_cost = sewing_cost_jacket + sewing_cost_trousers + sewing_cost_suit
    sewing_cost_label.set(f"Sewing Cost for jacket: {sewing_cost_jacket}\n"
                          f"Sewing Cost for trousers: {sewing_cost_trousers}\n"
                          f"Sewing Cost for suit: {sewing_cost_suit}\n"
                          f"Total Sewing Cost: {total_sewing_cost}")

layout = [
    [size_input, fabric_cost_input],
    [fabric_label],
    [sewing_cost_label, guietta.Button('Calculate', action=calculate)],
    [guietta.Button('Save to .docx', action=save_to_docx), guietta.Button('Save to .xlsx', action=save_to_xls)]
]

app.setLayout(layout)
app.run()