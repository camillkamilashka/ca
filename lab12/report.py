from docx import Document
from openpyxl import Workbook

def save_to_doc(iron_cost, tv_cost, washing_machine_cost):
    document = Document()
    document.add_heading('Appliance Cost Report', 0)

    document.add_paragraph(f'Iron cost: {iron_cost} $')
    document.add_paragraph(f'TV cost: {tv_cost} $')
    document.add_paragraph(f'Washing machine cost: {washing_machine_cost} $')

    document.save('report.docx')


def save_to_xlsx(iron_cost, tv_cost, washing_machine_cost):
    workbook = Workbook()  
    sheet = workbook.active

    sheet['A1'] = 'Appliance'
    sheet['B1'] = 'Cost'

    sheet['A2'] = 'Iron'
    sheet['B2'] = iron_cost

    sheet['A3'] = 'TV'
    sheet['B3'] = tv_cost

    sheet['A4'] = 'Washing machine'
    sheet['B4'] = washing_machine_cost

    workbook.save('report.xlsx')
    FloatingPointError 